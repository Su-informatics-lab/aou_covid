#!/usr/bin/env python3
"""
08_build_maintext_tables.py — Tables 1-3 in one consistent house style.

Table 1 is generated from the pipeline CSVs (see 07_build_tables_docx.py).
Tables 2 and 3 are transcribed verbatim from manuscript v18.3 — no value is
changed — and only the *presentation* is brought in line with Table 1:

  - three horizontal rules only, no vertical rules, no shading
  - 7 pt Arial, bold group headers, sub-rows indented
  - "n (%)" and "AOR (95% CI)" collapsed into single right-aligned cells,
    matching Table 1's "2,441 (60.1)" pattern, so all three tables read alike

Label harmonisation applied (content fix, not cosmetic): v18.3 wrote the income
reference four different ways across Table 2, Table 3, Results and the Figure 5
caption ("$35,000-99,999 (ref)", "ref: $35-100K", "$35-100K reference",
"35,000 to $99,999") and the education reference two ways. Tables now use the
spelled-out form throughout, matching Table 1 and Figure 4.

Usage:  python 08_build_maintext_tables.py
Output: JAMIA_tables_1to3.docx
"""
import csv
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

RULE = dict(val="single", sz="8", color="333333", space="0")
NO_RULE = None
NA = "NA"
HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.environ.get("TBL_OUT", HERE)

# ───────────────────────────── shared formatting helpers ─────────────────────


def set_borders(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:tcBorders")
    for edge, props in kw.items():
        e = OxmlElement(f"w:{edge}")
        if props is None:
            e.set(qn("w:val"), "none")
        else:
            for k, v in props.items():
                e.set(qn(f"w:{k}"), str(v))
        el.append(e)
    tcPr.append(el)


def white(cell):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "FFFFFF")
    sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)


def put(cell, text, bold=False, align="right", indent=0, size=7):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(1.2), Pt(1.2)
    if indent:
        pf.left_indent = Pt(9 * indent)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0, 0, 0)


def rules(tbl, spanner_row=False):
    """Nature/JAMIA three-rule system: above header, below header, below body."""
    last = len(tbl.rows) - 1
    hdr_end = 1 if spanner_row else 0
    for i, row in enumerate(tbl.rows):
        for cell in row.cells:
            white(cell)
            if i == 0 and spanner_row:
                set_borders(cell, top=RULE, bottom=NO_RULE, start=NO_RULE, end=NO_RULE)
            elif i == hdr_end:
                set_borders(cell, top=(NO_RULE if spanner_row else RULE),
                            bottom=RULE, start=NO_RULE, end=NO_RULE)
            elif i == last:
                set_borders(cell, top=NO_RULE, bottom=RULE, start=NO_RULE, end=NO_RULE)
            else:
                set_borders(cell, top=NO_RULE, bottom=NO_RULE, start=NO_RULE, end=NO_RULE)


def para(doc, text, size, bold=False, after=None):
    p = doc.add_paragraph()
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    return r


def simple_table(doc, headers, rows, widths, spanner=None):
    """Build a table where `rows` = [(label, indent, *cells)]; indent 0 = group header."""
    n_head = 2 if spanner else 1
    tbl = doc.add_table(rows=n_head + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    if spanner:
        h0 = tbl.rows[0].cells
        put(h0[0], "", bold=True, align="left")
        for label, j0, j1 in spanner:
            merged = h0[j0].merge(h0[j1])
            put(merged, label, bold=True, align="center")

    hrow = tbl.rows[n_head - 1].cells
    for j, t in enumerate(headers):
        put(hrow[j], t, bold=True, align="left" if j == 0 else "right")

    for i, r in enumerate(rows):
        label, indent = r[0], r[1]
        vals = r[2:]
        cells = tbl.rows[i + n_head].cells
        is_group = (indent == 0 and all(v == "" for v in vals))
        put(cells[0], label, bold=is_group, align="left", indent=indent)
        for j, v in enumerate(vals):
            put(cells[j + 1], v, align="right")

    rules(tbl, spanner_row=bool(spanner))
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
    return tbl


# ───────────────────────────── Table 1 (from pipeline CSVs) ──────────────────

SHARED = [
    ("Sex", None, "Sex"),
    ("N", "  Female", "Female"),
    ("N", "  Male", "Male"),
    ("N", "  Other", "Other"),
    ("Age group", None, "Age group, years"),
    ("Age group", "  <45", "<45"),
    ("Age group", "  45-54", "45–54"),
    ("Age group", "  55-64", "55–64"),
    ("Age group", "  65+", "≥65"),
    ("Age group", "  Mean age (SD)", "Mean age (s.d.)"),
    ("Vaccination", None, None),
    ("Vaccination", "  Vaccinated", "Recorded before index"),
    ("Vaccination", "  Unknown", "No record"),
    ("Charlson comorbidities", None, None),
    ("Charlson comorbidities", "  Myocardial Infarction", "Myocardial infarction"),
    ("Charlson comorbidities", "  Congestive Heart Failure", "Congestive heart failure"),
    ("Charlson comorbidities", "  Peripheral Vascular Disease", "Peripheral vascular disease"),
    ("Charlson comorbidities", "  Cerebrovascular Disease", "Cerebrovascular disease"),
    ("Charlson comorbidities", "  Dementia", "Dementia"),
    ("Charlson comorbidities", "  Chronic Pulmonary Disease", "Chronic pulmonary disease"),
    ("Charlson comorbidities", "  Rheumatic Disease", "Rheumatic disease"),
    ("Charlson comorbidities", "  Peptic Ulcer Disease", "Peptic ulcer disease"),
    ("Charlson comorbidities", "  Liver Disease Mild", "Liver disease, mild"),
    ("Charlson comorbidities", "  Liver Disease Moderate Severe", "Liver disease, moderate/severe"),
    ("Charlson comorbidities", "  Diabetes without Chronic Complications", "Diabetes without complications"),
    ("Charlson comorbidities", "  Diabetes with Chronic Complications", "Diabetes with complications"),
    ("Charlson comorbidities", "  Hemiplegia Paraplegia", "Hemiplegia or paraplegia"),
    ("Charlson comorbidities", "  Renal Disease Mild Moderate", "Renal disease, mild/moderate"),
    ("Charlson comorbidities", "  Renal Disease Severe", "Renal disease, severe"),
    ("Charlson comorbidities", "  HIV", "HIV"),
    ("Charlson comorbidities", "  Metastatic Solid Tumor", "Metastatic solid tumor"),
    ("Charlson comorbidities", "  Malignancy", "Malignancy"),
    ("Charlson comorbidities", "  AIDS", "AIDS"),
]
AOU_ONLY = [("Race", ["  White", "  Black", "  Asian", "  Other"]),
            ("Ethnicity", ["  Not Hispanic", "  Hispanic", "  Other"])]


def load_pipeline(path):
    d, section = {}, None
    with open(path) as fh:
        for r in csv.DictReader(fh):
            v = r["Variable"]
            if not v.startswith("  "):
                section = v
            cn, cp = r["Cases_n"].strip(), r["Cases_pct"].strip()
            kn, kp = r["Controls_n"].strip(), r["Controls_pct"].strip()
            d[(section, v)] = (f"{cn} {cp}".strip() if cp else cn,
                               f"{kn} {kp}".strip() if kp else kn)
    return d


def build_table1(doc):
    aou = load_pipeline(os.path.join(HERE, "results/aou_v7/table1_demographics.csv"))
    ms = load_pipeline(os.path.join(HERE, "results/ms/table1_demographics.csv"))
    A, M = aou[("N", "N")], ms[("N", "N")]

    para(doc, "Table 1. Demographic and clinical characteristics of the matched "
              "All of Us and MarketScan cohorts.", 9, bold=True)
    para(doc, "Data are n (%) unless otherwise specified. Cases were hospitalized within "
              "14 days of the COVID-19 index date; controls were COVID-positive participants "
              "not hospitalized within that window, matched 1:4 on encounter-density proxies. "
              "Percentages are column percentages within each cohort. Abbreviations: s.d., "
              "standard deviation.", 7.5)

    plan = []
    for sec, lab, disp in SHARED:
        plan.append((sec, lab, disp if disp else sec, False))
        if sec == "N" and lab == "  Other":
            for grp, members in AOU_ONLY:
                plan.append((grp, None, grp, True))
                for m in members:
                    plan.append((grp, m, m.strip(), True))

    rows = [("N", 0, A[0], A[1], M[0], M[1])]
    for sec, lab, disp, aou_only in plan:
        if lab is None:
            rows.append((disp, 0, "", "", "", ""))
            continue
        ac, ak = aou.get((sec, lab), ("", ""))
        mc, mk = (NA, NA) if aou_only else ms.get((sec, lab), (NA, NA))
        rows.append((disp, 1, ac or NA, ak or NA, mc or NA, mk or NA))

    simple_table(
        doc,
        ["Characteristic", "Cases", "Controls", "Cases", "Controls"],
        rows,
        [Inches(2.05), Inches(1.02), Inches(1.14), Inches(1.14), Inches(1.20)],
        spanner=[("All of Us", 1, 2), ("MarketScanᶜ", 3, 4)],
    )
    for fn in [
        "ᵃNA, not available. Race and ethnicity are not captured in MarketScan commercial claims.",
        "ᵇVaccination “No record” denotes no vaccination record identified before the index "
        "date using the concept set in eTable 3; it does not confirm the participant was "
        "unvaccinated.",
        "ᶜMarketScan captures commercially insured individuals, so adults aged ≥65 years, who "
        "are predominantly covered by Medicare, are under-represented. Plan type and geographic "
        "region for the MarketScan cohort are reported in eTable 9.",
    ]:
        para(doc, fn, 6.5, after=0)


# ───────────────────────────── Table 2 (verbatim from v18.3) ─────────────────
# Labels harmonised to the spelled-out form used in Table 1 and Figure 4.
T2 = [
    ("Income", 0, "", ""),
    ("<$10,000", 1, "675 (16.6)", "2,257 (14.2)"),
    ("$10,000–24,999", 1, "650 (16.0)", "2,268 (14.3)"),
    ("$25,000–34,999", 1, "318 (7.8)", "1,351 (8.5)"),
    ("$35,000–99,999 (reference)", 1, "793 (19.5)", "4,264 (26.9)"),
    ("$100,000–149,999", 1, "277 (6.8)", "1,292 (8.1)"),
    ("$150,000–199,999", 1, "122 (3.0)", "582 (3.7)"),
    ("≥$200,000", 1, "143 (3.5)", "705 (4.4)"),
    ("Missing", 1, "1,086 (26.7)", "3,137 (19.8)"),
    ("Employment", 0, "", ""),
    ("Employed (reference)", 1, "1,099 (27.0)", "5,930 (37.4)"),
    ("Student", 1, "84 (2.1)", "318 (2.0)"),
    ("Unemployed", 1, "1,334 (32.8)", "4,407 (27.8)"),
    ("Retired or other", 1, "1,427 (35.1)", "4,819 (30.4)"),
    ("Missing", 1, "120 (3.0)", "382 (2.4)"),
    ("Education", 0, "", ""),
    ("Never attended school", 1, "<20", "<20"),
    ("Below high school or GED", 1, "565 (13.9)", "1,643 (10.4)"),
    ("High school, GED, or some college", 1, "2,128 (52.4)", "8,164 (51.5)"),
    ("College graduate or higher (reference)", 1, "1,243 (30.6)", "5,758 (36.3)"),
    ("Missing", 1, "112 (2.8)", "273 (1.7)"),
    ("Housing tenure", 0, "", ""),
    ("Own home (reference)", 1, "1,519 (37.4)", "6,816 (43.0)"),
    ("Rent", 1, "2,030 (50.0)", "6,849 (43.2)"),
    ("Other", 1, "345 (8.5)", "1,615 (10.2)"),
    ("Missing", 1, "169 (4.2)", "571 (3.6)"),
    ("Housing stability", 0, "", ""),
    ("Stable (reference)", 1, "3,281 (80.7)", "12,967 (81.8)"),
    ("Unstable", 1, "710 (17.5)", "2,684 (16.9)"),
    ("Missing", 1, "73 (1.8)", "205 (1.3)"),
    ("Disability (any)", 0, "", ""),
    ("Yes", 1, "557 (13.7)", "2,174 (13.7)"),
    ("No (reference)", 1, "815 (20.1)", "3,349 (21.1)"),
    ("Missing*", 1, "2,692 (66.2)", "10,333 (65.2)"),
    ("Insurance type", 0, "", ""),
    ("Employer (reference)", 1, "788 (19.4)", "4,201 (26.5)"),
    ("Medicare", 1, "940 (23.1)", "3,496 (22.0)"),
    ("Medicaid", 1, "1,410 (34.7)", "4,159 (26.2)"),
    ("Other or none", 1, "368 (9.1)", "1,316 (8.3)"),
    ("Missing*", 1, "558 (13.7)", "2,684 (16.9)"),
]


def build_table2(doc):
    para(doc, "Table 2. Social determinants of health characteristics, matched "
              "All of Us cohort.", 9, bold=True)
    para(doc, "Data are n (%). Percentages are column percentages of the matched cohort "
              "(4,064 cases; 15,856 controls).", 7.5)
    simple_table(doc, ["Variable", "Cases", "Controls"], T2,
                 [Inches(3.05), Inches(1.55), Inches(1.55)])
    para(doc, "“Missing” denotes a Skip or Prefer-not-to-answer response. Cell counts <20 "
              "are suppressed per All of Us data use policy.", 6.5, after=0)
    para(doc, "*Disability and insurance missingness is mainly non-administration: the ACS-6 "
              "items were added to The Basics survey on November 10, 2020, and the insurance "
              "item after enrollment began.", 6.5, after=0)


# ───────────────────────────── Table 3 (verbatim from v18.3) ─────────────────
T3 = [
    ("Income (reference: $35,000–99,999)", 0, "", ""),
    ("<$10,000", 1, "1.46*** (1.29–1.65)", "1.18* (1.02–1.36)"),
    ("$10,000–24,999", 1, "1.37*** (1.21–1.55)", "1.18* (1.04–1.35)"),
    ("$25,000–34,999", 1, "1.19* (1.02–1.38)", "1.12 (0.96–1.31)"),
    ("$100,000–149,999", 1, "1.20* (1.03–1.41)", "1.24** (1.06–1.46)"),
    ("$150,000–199,999", 1, "1.19 (0.96–1.48)", "1.26* (1.01–1.57)"),
    ("≥$200,000", 1, "1.10 (0.90–1.34)", "1.15 (0.93–1.41)"),
    ("Missing", 1, "1.65*** (1.48–1.84)", "1.46*** (1.29–1.64)"),
    ("Insurance (reference: Employer)", 0, "", ""),
    ("Medicare", 1, "1.06 (0.93–1.20)", "0.95 (0.83–1.09)"),
    ("Medicaid", 1, "1.59*** (1.43–1.77)", "1.33*** (1.16–1.51)"),
    ("Other or none", 1, "1.31*** (1.13–1.52)", "1.15 (0.98–1.35)"),
    ("Education (reference: College graduate or higher)", 0, "", ""),
    ("Never attended school", 1, "3.35*** (1.66–6.78)", "2.85** (1.40–5.82)"),
    ("Below GED", 1, "1.35*** (1.18–1.53)", "1.13 (0.98–1.30)"),
    ("GED or some college", 1, "1.11* (1.02–1.21)", "1.02 (0.93–1.12)"),
    ("Employment (reference: Employed)", 0, "", ""),
    ("Unemployed", 1, "1.41*** (1.28–1.56)", "1.23*** (1.09–1.38)"),
    ("Student", 1, "1.61*** (1.24–2.09)", "1.52** (1.16–1.98)"),
    ("Retired or other", 1, "1.36*** (1.22–1.51)", "1.26*** (1.12–1.42)"),
    ("Housing tenure (reference: Own home)", 0, "", ""),
    ("Rent", 1, "1.28*** (1.17–1.39)", "1.13* (1.03–1.25)"),
    ("Other", 1, "0.99 (0.86–1.13)", "0.86* (0.74–1.00)"),
    ("Housing stability (reference: Stable)", 0, "", ""),
    ("Unstable", 1, "1.01 (0.92–1.11)", "0.94 (0.85–1.04)"),
    ("Disability (reference: No)", 0, "", ""),
    ("Any disability", 1, "0.94 (0.83–1.07)", "0.86* (0.76–0.98)"),
]


def build_table3(doc):
    para(doc, "Table 3. Domain-specific and jointly adjusted associations of survey-derived "
              "SDoH with COVID-19 hospitalization.", 9, bold=True)
    para(doc, "Values are adjusted odds ratios (95% confidence intervals) from conditional "
              "logistic regression within matched strata. The domain-specific model adds one "
              "SDoH domain to the base model; the joint model adds all six domains "
              "simultaneously. Both adjust for sex, race, ethnicity, age, vaccination, "
              "pandemic wave, and 19 Charlson comorbidities.", 7.5)
    simple_table(doc, ["Variable", "Domain-specific AOR (95% CI)", "Joint AOR (95% CI)"], T3,
                 [Inches(2.45), Inches(1.85), Inches(1.85)])
    para(doc, "*P < 0.05; **P < 0.01; ***P < 0.001. Abbreviations: AOR, adjusted odds ratio; "
              "CI, confidence interval; GED, General Educational Development; SDoH, social "
              "determinants of health.", 6.5, after=0)


# ───────────────────────────── assemble ──────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(0.7)

build_table1(doc)
doc.add_page_break()
build_table2(doc)
doc.add_page_break()
build_table3(doc)

path = os.path.join(OUT, "JAMIA_tables_1to3.docx")
doc.save(path)
print("saved:", path)
print(f"  Table 1: {len(SHARED) + 8} rows | Table 2: {len(T2)} | Table 3: {len(T3)}")
