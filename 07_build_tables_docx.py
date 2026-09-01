#!/usr/bin/env python3
"""
07_build_tables_docx.py — build paste-ready .docx tables straight from the
pipeline CSVs, so no number is ever hand-transcribed.

Outputs:
  Table1_two_cohort.docx   AoU + MarketScan demographics/clinical (main text)
  eTable9_marketscan.docx  MarketScan matched cohort alone (supplement)

Reads:
  results/aou_v7/table1_demographics.csv
  results/ms/table1_demographics.csv

Style: Nature table conventions (three horizontal rules, no vertical rules,
no shading, 7 pt Arial, numerics right-aligned). Two deliberate deviations to
match this manuscript: title is "Table N." not "Table N |", and unavailable
cells read "NA" rather than an em dash (house style bans em dashes).

Usage:  python 07_build_tables_docx.py
"""
import csv
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

RULE = dict(val="single", sz="8", color="333333", space="0")
NO_RULE = None
NA = "NA"
HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.environ.get("TBL_OUT", HERE)

# rows shared between the two cohorts, in main-text order.
# key = label as it appears in the pipeline CSV "Variable" column
# (section, csv_label, display).  section is REQUIRED because "  Unknown" occurs
# under Vaccination, Plan type and Region, and "  Other" under Sex, Race and
# Ethnicity — keying on the label alone silently returns the wrong row.
SHARED = [
    # NOTE: 04_tables.py emits no "Sex" header row, so the sex rows sit under
    # the "N" section in the CSV. The "Sex" row below is presentational only.
    ("Sex", None, "Sex"),
    ("N", "  Female", "Female"),
    ("N", "  Male", "Male"),
    ("N", "  Other", "Other"),
    ("Age group", None, "Age group, years"),
    ("Age group", "  <45", "<45"),
    ("Age group", "  45-54", "45–54"),
    ("Age group", "  55-64", "55–64"),
    ("Age group", "  65+", "≥65"),
    ("Age group", "  Mean age (SD)", "Mean age (s.d.)"),
    ("Vaccination", None, None),
    ("Vaccination", "  Vaccinated", "Recorded before index"),
    ("Vaccination", "  Unknown", "No record"),
    ("Charlson comorbidities", None, None),
    ("Charlson comorbidities", "  Myocardial Infarction", "Myocardial infarction"),
    ("Charlson comorbidities", "  Congestive Heart Failure", "Congestive heart failure"),
    ("Charlson comorbidities", "  Peripheral Vascular Disease", "Peripheral vascular disease"),
    ("Charlson comorbidities", "  Cerebrovascular Disease", "Cerebrovascular disease"),
    ("Charlson comorbidities", "  Dementia", "Dementia"),
    ("Charlson comorbidities", "  Chronic Pulmonary Disease", "Chronic pulmonary disease"),
    ("Charlson comorbidities", "  Rheumatic Disease", "Rheumatic disease"),
    ("Charlson comorbidities", "  Peptic Ulcer Disease", "Peptic ulcer disease"),
    ("Charlson comorbidities", "  Liver Disease Mild", "Liver disease, mild"),
    ("Charlson comorbidities", "  Liver Disease Moderate Severe", "Liver disease, moderate/severe"),
    ("Charlson comorbidities", "  Diabetes without Chronic Complications", "Diabetes without complications"),
    ("Charlson comorbidities", "  Diabetes with Chronic Complications", "Diabetes with complications"),
    ("Charlson comorbidities", "  Hemiplegia Paraplegia", "Hemiplegia or paraplegia"),
    ("Charlson comorbidities", "  Renal Disease Mild Moderate", "Renal disease, mild/moderate"),
    ("Charlson comorbidities", "  Renal Disease Severe", "Renal disease, severe"),
    ("Charlson comorbidities", "  HIV", "HIV"),
    ("Charlson comorbidities", "  Metastatic Solid Tumor", "Metastatic solid tumor"),
    ("Charlson comorbidities", "  Malignancy", "Malignancy"),
    ("Charlson comorbidities", "  AIDS", "AIDS"),
]
# AoU-only blocks, inserted after the Sex block
AOU_ONLY = [
    ("Race", ["  White", "  Black", "  Asian", "  Other"]),
    ("Ethnicity", ["  Not Hispanic", "  Hispanic", "  Other"]),
]
# MarketScan-only blocks: these are what eTable 9 keeps
MS_ONLY = {"Plan type", "Region"}


def load(path):
    """Return ({(section, label): (cases, controls)}, [(section, label), ...]).

    Keyed on (section, label) so duplicate labels cannot collide.
    """
    d, order, section = {}, [], None
    with open(path) as fh:
        for r in csv.DictReader(fh):
            v = r["Variable"]
            if not v.startswith("  "):
                section = v
            cn, cp = r["Cases_n"].strip(), r["Cases_pct"].strip()
            kn, kp = r["Controls_n"].strip(), r["Controls_pct"].strip()
            cs = f"{cn} {cp}".strip() if cp else cn
            ks = f"{kn} {kp}".strip() if kp else kn
            key = (section, v)
            if key in d:
                raise ValueError(f"duplicate row {key} in {path}")
            d[key] = (cs, ks)
            order.append(key)
    return d, order


def set_borders(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:tcBorders")
    for edge, props in kw.items():
        e = OxmlElement(f"w:{edge}")
        if props is None:
            e.set(qn("w:val"), "none")
        else:
            for k, v in props.items():
                e.set(qn(f"w:{k}"), str(v))
        el.append(e)
    tcPr.append(el)


def white(cell):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "FFFFFF")
    sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)


def put(cell, text, bold=False, align="right", indent=0, size=7):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(1.2), Pt(1.2)
    if indent:
        pf.left_indent = Pt(9 * indent)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0, 0, 0)


def rules(tbl, spanner_row=False):
    last = len(tbl.rows) - 1
    hdr_end = 1 if spanner_row else 0
    for i, row in enumerate(tbl.rows):
        for cell in row.cells:
            white(cell)
            if i == 0 and spanner_row:
                set_borders(cell, top=RULE, bottom=NO_RULE, start=NO_RULE, end=NO_RULE)
            elif i == hdr_end:
                set_borders(cell, top=(RULE if not spanner_row else NO_RULE),
                            bottom=RULE, start=NO_RULE, end=NO_RULE)
            elif i == last:
                set_borders(cell, top=NO_RULE, bottom=RULE, start=NO_RULE, end=NO_RULE)
            else:
                set_borders(cell, top=NO_RULE, bottom=NO_RULE, start=NO_RULE, end=NO_RULE)


def para(doc, text, size, bold=False, after=None):
    p = doc.add_paragraph()
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    return p


# ==================================================== load both cohorts
aou, aou_order = load(os.path.join(HERE, "results/aou_v7/table1_demographics.csv"))
ms, ms_order = load(os.path.join(HERE, "results/ms/table1_demographics.csv"))

AOU_N = aou[("N", "N")]
MS_N = ms[("N", "N")]
print("AoU  N:", AOU_N)
print("MS   N:", MS_N)

# ==================================================== Table 1 (two cohort)
doc = Document()
s = doc.sections[0]
s.left_margin = s.right_margin = Inches(0.7)

para(doc, "Table 1. Demographic and clinical characteristics of the matched "
          "All of Us and MarketScan cohorts.", 9, bold=True)
para(doc, "Data are n (%) unless otherwise specified. Cases were hospitalized within 14 days "
          "of the COVID-19 index date; controls were COVID-positive participants not hospitalized "
          "within that window, matched 1:4 on encounter-density proxies. Percentages are column "
          "percentages within each cohort. Abbreviations: s.d., standard deviation.", 7.5)

# assemble row plan: Sex, [Race, Ethnicity], Age, Vaccination, Charlson.
# plan entries are (section, csv_label_or_None, display, aou_only)
plan = []
for sec, lab, disp in SHARED:
    plan.append((sec, lab, disp if disp else sec, False))
    if sec == "N" and lab == "  Other":         # insert AoU-only blocks after Sex
        for grp, members in AOU_ONLY:
            plan.append((grp, None, grp, True))
            for m in members:
                plan.append((grp, m, m.strip(), True))

tbl = doc.add_table(rows=2 + 1 + len(plan), cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.autofit = False

h0 = tbl.rows[0].cells
a = h0[1].merge(h0[2])
b = h0[3].merge(h0[4])
put(h0[0], "", bold=True, align="left")
put(a, "All of Us", bold=True, align="center")
put(b, "MarketScanᶜ", bold=True, align="center")

for j, t in enumerate(["Characteristic", "Cases", "Controls", "Cases", "Controls"]):
    put(tbl.rows[1].cells[j], t, bold=True, align="left" if j == 0 else "right")

put(tbl.rows[2].cells[0], "N", align="left")
for j, v in enumerate([AOU_N[0], AOU_N[1], MS_N[0], MS_N[1]]):
    put(tbl.rows[2].cells[j + 1], v, align="right")

for i, (sec, lab, disp, aou_only) in enumerate(plan):
    cells = tbl.rows[i + 3].cells
    is_group = lab is None
    put(cells[0], disp, bold=is_group, align="left", indent=0 if is_group else 1)
    if is_group:
        for j in range(4):
            put(cells[j + 1], "", align="right")
        continue
    ac, ak = aou.get((sec, lab), ("", ""))
    mc, mk = (NA, NA) if aou_only else ms.get((sec, lab), (NA, NA))
    for j, v in enumerate([ac, ak, mc, mk]):
        put(cells[j + 1], v if v else NA, align="right")

rules(tbl, spanner_row=True)
for cell in (tbl.rows[0].cells[1], tbl.rows[0].cells[3]):
    set_borders(cell, top=RULE, bottom=RULE, start=NO_RULE, end=NO_RULE)
W = [Inches(2.05), Inches(1.02), Inches(1.14), Inches(1.14), Inches(1.20)]
for row in tbl.rows:
    for j, cell in enumerate(row.cells):
        cell.width = W[j]

for fn in [
    "ᵃNA, not available. Race and ethnicity are not captured in MarketScan commercial claims.",
    "ᵇVaccination “No record” denotes no vaccination record identified before the "
    "index date using the concept set in eTable 3; it does not confirm the participant was unvaccinated.",
    "ᶜMarketScan captures commercially insured individuals, so adults aged ≥65 years, "
    "who are predominantly covered by Medicare, are under-represented. Plan type, region, and "
    "pandemic wave for the MarketScan cohort are reported in eTable 9.",
]:
    para(doc, fn, 6.5, after=0)

p1 = os.path.join(OUT, "Table1_two_cohort.docx")
doc.save(p1)
print("saved:", p1)

# ==================================================== eTable 9 (MarketScan)
doc2 = Document()
s = doc2.sections[0]
s.left_margin = s.right_margin = Inches(0.9)

para(doc2, "eTable 9. Insurance plan type and geographic region, MarketScan matched cohort.",
     9, bold=True)
para(doc2, "Data are n (%). Counts are from the post-matching, post-cutoff-trim analytic "
           "cohort. Demographic, vaccination, and comorbidity characteristics for this cohort "
           "are reported alongside All of Us in Table 1 and are not repeated here. Plan type "
           "and region are site-specific covariates in the MarketScan model. Cells with fewer "
           "than 20 participants are suppressed.", 7.5)

# Keep ONLY the MarketScan-specific blocks. Everything else now lives in Table 1,
# so each number has exactly one home and the two tables cannot drift apart.
# Section-tracked because "  Unknown" occurs under Vaccination, Plan type AND Region.
rows2 = [(sec, lab) for (sec, lab) in ms_order
         if lab == "N" or sec in MS_ONLY]
t2 = doc2.add_table(rows=1 + len(rows2), cols=3)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2.autofit = False
for j, t in enumerate(["Characteristic", "Cases", "Controls"]):
    put(t2.rows[0].cells[j], t, bold=True, align="left" if j == 0 else "right")
for i, key in enumerate(rows2):
    cells = t2.rows[i + 1].cells
    cs, ks = ms[key]
    label = key[1]
    is_group = (cs == "" and ks == "")
    put(cells[0], label.strip(), bold=is_group, align="left",
        indent=0 if (is_group or label == "N") else 1)
    put(cells[1], cs, align="right")
    put(cells[2], ks, align="right")
rules(t2, spanner_row=False)
W2 = [Inches(3.0), Inches(1.7), Inches(1.7)]
for row in t2.rows:
    for j, cell in enumerate(row.cells):
        cell.width = W2[j]

para(doc2, "Abbreviations: CDHP, consumer-driven health plan; EPO, exclusive provider "
           "organization; HDHP, high-deductible health plan; HMO, health maintenance "
           "organization; POS, point of service; PPO, preferred provider organization.",
     6.5, after=0)

p2 = os.path.join(OUT, "eTable9_marketscan.docx")
doc2.save(p2)
print("saved:", p2)
